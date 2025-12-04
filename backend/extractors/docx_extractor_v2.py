"""
Enhanced DOCX extraction module with improved formatting preservation.
Extracts text with bold, italic, lists, and better table handling.
"""
import logging
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

logger = logging.getLogger("pinguin_backend.extractors.docx_v2")


class EnhancedDOCXExtractor:
    """
    Enhanced DOCX extractor with formatting preservation.
    
    Features:
    - Bold and italic text preservation (using markdown)
    - List detection (bullets and numbered)
    - Improved table formatting with header detection
    - Heading hierarchy with proper markdown
    - Code block detection
    """
    
    def __init__(self):
        """Initialize enhanced DOCX extractor."""
        self.logger = logger
    
    def extract(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from a DOCX file with formatting.
        
        Args:
            file_path: Path to the DOCX file
        
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is invalid
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not path.suffix.lower() in ['.docx', '.doc']:
            raise ValueError(f"File is not a Word document: {file_path}")
        
        self.logger.info(f"Extracting text from DOCX with formatting: {path.name}")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            self.logger.error(f"Error opening DOCX file: {e}")
            raise ValueError(f"Failed to open DOCX file: {str(e)}")
        
        # Extract text with structure and formatting
        text_parts = []
        headings = []
        paragraph_count = 0
        table_count = 0
        list_count = 0
        
        # Process document elements in order
        for element in doc.element.body:
            # Check if it's a paragraph
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                para_text, para_info = self._process_paragraph_with_formatting(para)
                
                if para_text:
                    if para_info['is_heading']:
                        # Add heading with markdown
                        heading_marker = '#' * para_info['heading_level']
                        text_parts.append(f"\n{heading_marker} {para_text}\n")
                        headings.append({
                            'level': para_info['heading_level'],
                            'text': para_text,
                            'position': len(text_parts)
                        })
                    elif para_info['is_list_item']:
                        # Add list item with marker
                        marker = para_info['list_marker']
                        text_parts.append(f"{marker} {para_text}")
                        list_count += 1
                    else:
                        text_parts.append(para_text)
                        paragraph_count += 1
            
            # Check if it's a table
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_text = self._process_table_with_formatting(table)
                if table_text:
                    text_parts.append(f"\n[Table {table_count + 1}]\n{table_text}\n")
                    table_count += 1
        
        full_text = "\n".join(text_parts)
        
        metadata = {
            'filename': path.name,
            'file_type': 'docx',
            'extractor': 'enhanced_python-docx',
            'paragraph_count': paragraph_count,
            'heading_count': len(headings),
            'table_count': table_count,
            'list_count': list_count,
            'headings': headings,
            'char_count': len(full_text),
            'has_formatting': True
        }
        
        self.logger.info(
            f"Extracted {len(full_text)} characters. "
            f"Paragraphs: {paragraph_count}, "
            f"Headings: {len(headings)}, "
            f"Tables: {table_count}, "
            f"Lists: {list_count}"
        )
        
        return full_text, metadata
    
    def _process_paragraph_with_formatting(self, paragraph: Paragraph) -> Tuple[str, Dict]:
        """
        Process a paragraph with formatting preservation.
        
        Args:
            paragraph: Paragraph object from python-docx
            
        Returns:
            Tuple of (formatted_text, paragraph_info)
        """
        if not paragraph.text.strip():
            return "", {'is_heading': False, 'is_list_item': False}
        
        # Check if paragraph is a heading
        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = style_name.startswith('Heading')
        
        heading_level = 0
        if is_heading:
            try:
                heading_level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                heading_level = 1
        
        # Check if paragraph is a list item
        is_list_item = style_name.startswith('List') or paragraph.text.strip().startswith(('•', '-', '*', '○'))
        list_marker = '-'
        
        if is_list_item:
            # Determine list type
            if any(paragraph.text.strip().startswith(m) for m in ['•', '○', '-', '*']):
                list_marker = '-'
            else:
                # Check for numbered list
                first_word = paragraph.text.strip().split()[0] if paragraph.text.strip() else ''
                if first_word.rstrip('.').isdigit():
                    list_marker = first_word
        
        # Extract text with inline formatting
        formatted_text = self._extract_formatted_text(paragraph)
        
        para_info = {
            'is_heading': is_heading,
            'heading_level': heading_level,
            'is_list_item': is_list_item,
            'list_marker': list_marker
        }
        
        return formatted_text, para_info
    
    def _extract_formatted_text(self, paragraph: Paragraph) -> str:
        """
        Extract text with inline formatting (bold, italic) using markdown.
        
        Args:
            paragraph: Paragraph object
            
        Returns:
            Formatted text string with markdown
        """
        text_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            # Check for code (monospace font)
            if run.font.name and 'mono' in run.font.name.lower():
                text = f"`{text}`"
            
            text_parts.append(text)
        
        return ''.join(text_parts).strip()
    
    def _process_table_with_formatting(self, table: Table) -> str:
        """
        Process a table with improved formatting and header detection.
        
        Args:
            table: Table object from python-docx
            
        Returns:
            Markdown formatted table string
        """
        if not table.rows:
            return ""
        
        table_data = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Extract cell text with formatting
                cell_text = []
                for para in cell.paragraphs:
                    formatted = self._extract_formatted_text(para)
                    if formatted:
                        cell_text.append(formatted)
                
                row_cells.append(' '.join(cell_text).strip())
            
            table_data.append(row_cells)
        
        if not table_data:
            return ""
        
        # Detect header row (first row is usually header)
        has_header = len(table_data) > 1
        
        # Check if first row has bold text (common for headers)
        if has_header and table.rows[0].cells:
            first_row_bold = any(
                run.bold 
                for cell in table.rows[0].cells 
                for para in cell.paragraphs 
                for run in para.runs
            )
            has_header = first_row_bold or has_header
        
        # Format as markdown table
        return self._format_table_as_markdown(table_data, has_header)
    
    def _format_table_as_markdown(self, table_data: List[List[str]], has_header: bool = True) -> str:
        """
        Format table as markdown with proper alignment.
        
        Args:
            table_data: Table data as list of lists
            has_header: Whether first row is header
            
        Returns:
            Markdown formatted table string
        """
        if not table_data:
            return ""
        
        lines = []
        
        # Determine column widths
        col_count = max(len(row) for row in table_data)
        col_widths = [0] * col_count
        
        for row in table_data:
            for i, cell in enumerate(row):
                if i < col_count:
                    # Remove markdown formatting for width calculation
                    clean_cell = cell.replace('**', '').replace('*', '').replace('`', '')
                    col_widths[i] = max(col_widths[i], len(clean_cell))
        
        # Format rows
        for row_idx, row in enumerate(table_data):
            # Pad row to col_count
            padded_row = list(row) + [''] * (col_count - len(row))
            
            # Format cells
            formatted_cells = []
            for i, cell in enumerate(padded_row):
                # Calculate padding (accounting for markdown)
                clean_cell = cell.replace('**', '').replace('*', '').replace('`', '')
                padding = col_widths[i] - len(clean_cell)
                formatted_cells.append(cell + ' ' * padding)
            
            lines.append('| ' + ' | '.join(formatted_cells) + ' |')
            
            # Add separator after header
            if has_header and row_idx == 0:
                separator = '|' + '|'.join(['-' * (w + 2) for w in col_widths]) + '|'
                lines.append(separator)
        
        return '\n'.join(lines)
    
    def get_heading_hierarchy(self, file_path: str) -> List[Dict]:
        """
        Extract heading hierarchy from document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of heading dictionaries with level and text
        """
        try:
            doc = Document(file_path)
            headings = []
            
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                if style_name.startswith('Heading'):
                    try:
                        level = int(style_name.split()[-1])
                        text = para.text.strip()
                        if text:
                            headings.append({
                                'level': level,
                                'text': text
                            })
                    except (ValueError, IndexError):
                        continue
            
            return headings
        
        except Exception as e:
            self.logger.error(f"Error extracting heading hierarchy: {e}")
            return []
